from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE = Path(__file__).resolve().parents[1]
OUT = BASE / "ADMIN_MANUAL.docx"
SS = BASE / "manual_screenshots"

MAROON = "7A0C2E"
GOLD = "F2A900"
INK = "1E2330"
MUTED = "5B6270"
LIGHT = "F6F7FA"
LINE = "D9DEE8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
      set_cell_text(hdr[i], h, bold=True, color=MAROON)
      set_cell_shading(hdr[i], "FFF4D7")
      if widths:
          hdr[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val))
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(MAROON if level <= 2 else INK)
    return p


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF8E6")
    p = cell.paragraphs[0]
    r = p.add_run(title + " ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(MAROON)
    r.font.size = Pt(10)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_image(doc, file_name, caption, width=5.9):
    path = SS / file_name
    if not path.exists():
        add_note(doc, "Screenshot missing:", str(path))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size in [("Heading 1", 17), ("Heading 2", 13.5), ("Heading 3", 11.5)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(MAROON)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)

    for name in ["List Bullet", "List Number"]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(10.2)


doc = Document()
configure_document(doc)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("TransitAI UTM\nAdministrator Manual")
r.font.name = "Arial"
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(MAROON)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("Web-Based Campus Transport Chatbot Prototype\nMECS0033-52 Artificial Intelligence · Group 5")
sr.font.name = "Arial"
sr.font.size = Pt(11)
sr.font.color.rgb = RGBColor.from_string(MUTED)

add_note(
    doc,
    "Submission purpose:",
    "This Word manual is the teacher-facing administrator guide. It documents the prototype design, verified and simulated data boundaries, usage process, configuration notes, and screenshots for each required interactive flow."
)

add_heading(doc, "1. Prototype Scope", 1)
doc.add_paragraph(
    "TransitAI UTM is a self-contained browser prototype. It demonstrates how an AI campus transport assistant can receive student input, classify intent, retrieve relevant shuttle knowledge, generate a grounded answer, and trigger delay notifications through a logic proof."
)
add_table(
    doc,
    ["Item", "Prototype decision"],
    [
        ["Platform", "Single-page HTML/CSS/JavaScript app; opens directly in a browser."],
        ["Backend", "None. No Firebase, API key, server, or database is required for grading."],
        ["AI concepts shown", "Intent recognition, RAG-style retrieval, conversation memory, predicate logic, resolution refutation."],
        ["Data policy", "Use public route labels/sequences where available. Mark missing or live data as prototype simulation."],
        ["Presentation safety", "Default demo clock is 10:00 so the app still shows active service during a night presentation."],
    ],
    widths=[Inches(1.7), Inches(4.8)],
)

add_heading(doc, "2. Data Source And Accuracy Policy", 1)
doc.add_paragraph(
    "The prototype avoids pretending that simulated data is official. The app and manual separate route information that could be aligned to public UTM-related sources from data that must remain mock/simulated."
)
add_table(
    doc,
    ["Data shown", "Status", "Administrator explanation"],
    [
        ["BAS A/B/C/D/E/F/G/H route labels and public stop sequences", "Public route listing aligned", "Used from UTM/KDOJ public route listings where available."],
        ["Stop labels such as CP, KP, K9/K10, KTC, KDOJ, KDSE, KTDI, N24, SKT, P19, T02, T08, V01", "Public stop labels", "Used because they appear in public route listings."],
        ["FKE query support", "Partially verified mapping", "Mapped to P19 / FKE area. The public route list uses P19; public bus-stop pages also show FKE-area stops."],
        ["FC query support", "Prototype mapping", "Mapped to N24 / Cluster Area because N24 appears in public route listings; exact FC stop mapping was not found."],
        ["PSZ Library", "Verified campus landmark, route stop unverified", "The app can describe PSZ but does not invent a direct route to PSZ."],
        ["Next departure, headway, operating hours, ETA", "Prototype simulation", "Used only to demonstrate response flow and demo-time behavior; not official timetable data."],
        ["Delay status", "Prototype simulation", "Staff Demo toggle simulates a staff-issued disruption so the resolution proof can be shown."],
    ],
)
add_bullets(doc, [
    "Source 1: UTM DVC Development shuttle bus schedule page - https://dvcdev.utm.my/announcement/shuttle-bus-schedule/",
    "Source 2: KDOJ UTM Bus Schedule route list - https://www.kdoj.com.my/insight/utm-bus-schedule/",
    "Source 3: UTM JB 2025 shuttle timetable PDF - https://studentppi.utmspace.edu.my/wp-content/uploads/2024/12/Jadual-Bas-Shuttle-Kampus-UTM-JB-2025-01012025.pdf",
])

add_heading(doc, "3. Component And Function Map", 1)
add_table(
    doc,
    ["File", "Main object", "Important functions", "Responsibility"],
    [
        ["js/kb.js", "KB, KBUtil", "findStop, routesBetween, nextDeparture, setClock, setDelayed", "Knowledge base, route data, demo clock, delay state."],
        ["js/intent.js", "Intent", "classify(query)", "Scores six intents and exposes confidence to the UI."],
        ["js/retriever.js", "Retriever", "retrieve(query, k)", "RAG-style token-overlap retrieval with source notes."],
        ["js/resolution.js", "Resolution", "prove(user, route)", "Resolution-refutation proof for delay notifications."],
        ["js/chat.js", "Chat", "send(text), init()", "Conversation controller, response cards, short-term memory."],
        ["js/app.js", "App", "showScreen, renderAdmin, renderAlerts", "Screen routing, Staff Demo controls, alerts, feedback, proof display."],
    ],
)
doc.add_paragraph("Runtime data flow:")
add_steps(doc, [
    "User enters a query or taps a quick action.",
    "Intent.classify() assigns an intent and confidence score.",
    "Retriever.retrieve() selects top knowledge chunks and a source note.",
    "The relevant handler renders a grounded response card.",
    "The side panel shows intent, retrieval, generation, and data source for transparency.",
])

add_heading(doc, "4. Student-Facing Usage With Screenshots", 1)
add_heading(doc, "4.1 Home / Chat", 2)
doc.add_paragraph("The first screen is already the usable chatbot, not a landing page. Students can type or use quick actions.")
add_image(doc, "01-home-chat.png", "Home screen with phone status time, quick actions, chat composer, and AI pipeline panel.")

add_heading(doc, "4.2 Schedule Inquiry", 2)
add_steps(doc, [
    "Tap Next bus or type: When is the next bus to FKE?",
    "Confirm the intent chip shows Schedule inquiry and a confidence score.",
    "Read the card: route, next departure, operating hours, demo time basis, and data source.",
])
add_image(doc, "02-schedule-demo-time.png", "Schedule answer using 10:00 demo time. Route sequence is public-list aligned; timing is simulated.")

add_heading(doc, "4.3 Route Guidance", 2)
add_steps(doc, [
    "Type or tap: How do I get from KTDI to P19 FKE?",
    "The chatbot recommends BAS G1/G2/G3 and displays the ordered stop sequence.",
    "Use the source note to explain that the route sequence is public-list aligned while timings remain simulated.",
])
add_image(doc, "03-route-guidance.png", "Route guidance card with BAS G stop sequence and source note.")

add_heading(doc, "4.4 Arrival ETA", 2)
add_steps(doc, [
    "Type: When will the next bus arrive at CP?",
    "The card shows simulated ETA, scheduled time, demo time basis, and fallback wording.",
    "Explain that production would replace this with live GPS or official timetable feed.",
])
add_image(doc, "04-arrival-simulated.png", "Arrival screen clearly labelled as simulated ETA.")

add_heading(doc, "4.5 Bus Stop Detail", 2)
add_steps(doc, [
    "Type: Where is the FC bus stop?",
    "The app opens the bus-stop detail screen.",
    "The data-status line tells the administrator whether this is public stop data, verified landmark data, or a prototype mapping.",
])
add_image(doc, "05-bus-stop-detail.png", "Bus stop detail screen with route serving data, status, facilities, and walking note.")

add_heading(doc, "5. Staff Demo Operation", 1)
add_heading(doc, "5.1 Demo Time Control", 2)
doc.add_paragraph(
    "The real current time caused a presentation problem because nighttime demos would show routes as closed. The Staff Demo panel now includes a fixed demo clock. Use 10:00 for normal demo, 20:30 to intentionally demonstrate closed-service behavior, or Live device time if required."
)
add_steps(doc, [
    "Open Staff Demo.",
    "Choose 08:00, 10:00, 13:00, 17:30, 20:30, or Live device time.",
    "Return to Chat and run a schedule or arrival query.",
    "Point out the Time basis row in the response card.",
])
add_image(doc, "06-admin-demo-time-control.png", "Staff Demo time selector plus public-route delay controls.")

add_heading(doc, "5.2 Delay Alert And Resolution-Refutation Proof", 2)
add_steps(doc, [
    "Open Staff Demo.",
    "Toggle BAS A1/A2 - KP to Lingkaran Ilmu to delayed.",
    "The app runs Resolution.prove('ali', 'route_a').",
    "The proof panel prints P1, P2, P3, P4, the negated goal, R1, R2, R3, and the empty clause.",
    "Open Alerts to show the resulting notification.",
])
add_image(doc, "07-resolution-proof.png", "Staff Demo proof view after triggering route_a delay.")
add_image(doc, "07b-resolution-proof-box.png", "Resolution proof trace: empty clause proves NotifyUser(ali, route_a, delay).", width=3.2)
add_image(doc, "08-alerts-after-proof.png", "Alerts screen after a successful proof-triggered delay notification.")

add_heading(doc, "5.3 Feedback / Escalation", 2)
add_steps(doc, [
    "Open Feedback.",
    "Choose issue type and enter details.",
    "Submit the report.",
    "The log records the issue for staff follow-up in the prototype session.",
])
add_image(doc, "09-feedback-log.png", "Feedback log after a sample report is submitted.")

add_heading(doc, "6. Configuration Guide", 1)
add_heading(doc, "6.1 Add Or Edit A Stop", 2)
doc.add_paragraph("Edit KB.stops in js/kb.js. Add aliases users might type. Add a status value so the UI can tell the truth about public vs simulated data.")
doc.add_paragraph(
    '{ id: "p19", name: "P19 / FKE Area", aliases: ["p19", "fke", "electrical"], status: "public stop/area label with faculty alias", landmark: "...", facilities: ["Faculty area"], walking: "..." }'
)
add_heading(doc, "6.2 Add Or Edit A Route", 2)
doc.add_paragraph("Edit KB.routes in js/kb.js. Keep public route sequences separate from simulated timing values.")
doc.add_paragraph(
    '{ id: "route_g", name: "BAS G1/G2/G3 - KTR/KTHO/KTDI to SKT", stops: ["ktr","ktho","ktdi","n24","skt","p19","cp"], operating: { start: "07:00", end: "18:30" }, headway: 20, source: "KDOJ public route list; timing simulated." }'
)
add_heading(doc, "6.3 Tune Intent Recognition", 2)
add_bullets(doc, [
    "Edit CUES in js/intent.js.",
    "Strong cues score 2; weak cues score 1.",
    "Avoid ultra-common words such as bare 'to' and 'from' as strong cues.",
    "After edits, test schedule, route guidance, arrival, bus stop, alerts, feedback, and FAQ queries.",
])
add_heading(doc, "6.4 Change Default Demo Time", 2)
doc.add_paragraph("Edit KB.clock.demoTime in js/kb.js. The default is 10:00 because it is a reliable presentation time. The Staff Demo selector can override it without editing code.")

add_heading(doc, "7. Troubleshooting", 1)
add_table(
    doc,
    ["Symptom", "Likely cause", "Fix"],
    [
        ["Blank screen", "Script load issue or missing file", "Reload; ensure all js/*.js files are present."],
        ["Schedule says closed now", "Demo/live clock outside operating window", "Staff Demo -> Demo time -> 10:00 for normal presentation."],
        ["Route missing", "No verified direct route in KB", "Use CP as interchange or add a verified route sequence; do not invent official data."],
        ["Wrong intent", "Missing cue phrase", "Add the phrase to CUES in js/intent.js and retest."],
        ["Proof does not run", "Route is not delayed or user is not subscribed", "Toggle delay and ensure KB.subscriptions includes ali for route_a."],
        ["Edited data not visible", "Browser cache", "Save and hard-refresh the browser."],
    ],
)

add_heading(doc, "8. Rubric Mapping", 1)
add_table(
    doc,
    ["Rubric criterion", "Evidence in prototype"],
    [
        ["Originality / Interactive Screen", "Phone-frame app, quick actions, free text, intent confidence, RAG-style panel, proof trace, Staff Demo controls."],
        ["Problem Solving", "Schedule uncertainty, route confusion, arrival uncertainty, delay alerts, and feedback are all covered."],
        ["Admin Manual", "This Word manual includes component map, configuration notes, source policy, screenshots, procedures, and troubleshooting."],
    ],
)

add_heading(doc, "9. Submission Checklist", 1)
add_bullets(doc, [
    "Open prototype/index.html or run python3 -m http.server 8000 from the prototype folder.",
    "Use Staff Demo -> Demo time -> 10:00 before presenting schedule and arrival flows.",
    "State clearly that ETA/headway/timetable are simulated POC data.",
    "Use the route source note when asked whether the route data is official.",
    "Show Staff Demo -> delay toggle -> proof trace -> Alerts as the originality highlight.",
    "Submit ADMIN_MANUAL.docx and TransitAI_UTM_Prototype_Demo.pptx with the prototype folder.",
])

footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("TransitAI UTM Prototype · Administrator Manual · Group 5").font.size = Pt(8)

doc.save(OUT)
print(OUT)
