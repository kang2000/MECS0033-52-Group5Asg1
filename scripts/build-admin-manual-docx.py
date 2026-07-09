from html import escape as h
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parents[1]
DOCX_OUT = BASE / "ADMIN_MANUAL.docx"
HTML_OUT = BASE / "admin_manual.html"
MD_OUT = BASE / "ADMIN_MANUAL.md"
SS = BASE / "manual_screenshots_annotated"
PLAIN = "000000"


manual_events = []


def set_font(run, size=None, bold=False, italic=False, monospace=False):
    name = "Courier New" if monospace else "Arial"
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.color.rgb = RGBColor.from_string(PLAIN)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)


def set_cell_margins(cell, top=55, start=95, bottom=55, end=95):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=8.4):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor.from_string(PLAIN)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1

    for name, size in [("Heading 1", 15.5), ("Heading 2", 12.5), ("Heading 3", 11.2)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(PLAIN)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    for name in ["List Bullet", "List Number"]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(10)
        styles[name].font.color.rgb = RGBColor.from_string(PLAIN)


def add_heading(doc, text, level=1):
    manual_events.append(("heading", level, text))
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size={1: 15.5, 2: 12.5, 3: 11.2}.get(level, 10.5), bold=True)
    return p


def add_para(doc, text=""):
    manual_events.append(("p", text))
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=10.3)
    return p


def add_bullets(doc, items):
    manual_events.append(("bullets", items))
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r, size=10)


def add_steps(doc, items):
    manual_events.append(("steps", items))
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(3)
        marker = p.add_run(f"{idx}. ")
        set_font(marker, size=10, bold=True)
        r = p.add_run(item)
        set_font(r, size=10)


def add_code(doc, text):
    manual_events.append(("code", text))
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=8.3, monospace=True)


def add_table(doc, headers, rows, widths=None, size=8.2):
    manual_events.append(("table", headers, rows))
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=size)
        if widths:
            table.rows[0].cells[i].width = widths[i]

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=size)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def add_image(doc, file_name, caption, width=6.55):
    manual_events.append(("image", file_name, caption))
    path = SS / file_name
    if not path.exists():
        add_para(doc, f"Screenshot missing: {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_font(r, size=8.5, italic=True)


def build_docx():
    doc = Document()
    props = doc.core_properties
    props.title = "TransitAI UTM Administrator Manual"
    props.subject = "Campus transport chatbot prototype administrator manual"
    props.author = "MECS0033-52 Group 5"
    props.last_modified_by = "MECS0033-52 Group 5"
    props.comments = ""
    props.keywords = "TransitAI UTM, campus transport, administrator manual"
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TransitAI UTM\nAdministrator Manual")
    set_font(r, size=23, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Web-Based Campus Transport Chatbot Prototype\nMECS0033-52 Artificial Intelligence, Section 52, Group 5")
    set_font(sr, size=10.5)

    add_para(doc, "This is the separate Administrator Manual deliverable. The in-app Staff Demo screen is a control surface for demonstration and proof execution; it does not replace this Word manual.")
    add_table(
        doc,
        ["Manual item", "Description"],
        [
            ["Audience", "Lecturer, marker, transport administrator, or maintainer reviewing the prototype."],
            ["Prototype path", "MECS0033-52/Group5Asg1/prototype/index.html"],
            ["Run mode", "Open index.html directly, or run python3 -m http.server 8000 from the prototype folder."],
            ["Main grading evidence", "Interactive screen, problem-solving workflow, administrator process, screenshots, and resolution proof."],
            ["Data boundary", "Public route names/stop sequences are used where available; timetable, ETA, delay state and walking notes are configured prototype estimates."],
        ],
        widths=[Inches(1.55), Inches(5.05)],
    )

    add_heading(doc, "1. Purpose and Rubric Alignment", 1)
    add_para(doc, "TransitAI UTM turns the project report into an inspectable browser prototype. The marker can click through the phone-style app, ask transport questions, change demo time, create feedback reports, toggle staff delays, and inspect the AI reasoning evidence.")
    add_table(
        doc,
        ["Rubric criterion", "How the prototype supports full marks", "Where to verify"],
        [
            ["Originality / Interactive Screen", "Home navigation, chatbot input, intent confidence, RAG-style evidence, timetable search, full-day route timetable, subscription-filtered alerts, feedback escalation, Staff Control Panel, and visible AI pipeline.", "Figures 1-13; index.html"],
            ["Problem Solving", "Addresses common campus transport pain points: route planning, next bus timing, arrival estimate, stop lookup, service delay notification, and passenger issue reporting.", "Chat, Timetable, Alerts, Feedback"],
            ["Admin Manual", "Explains component map, truthful data policy, configuration, operating steps, proof procedure, troubleshooting, and screenshot-guided usage.", "This ADMIN_MANUAL.docx"],
        ],
        widths=[Inches(1.55), Inches(3.7), Inches(1.35)],
    )

    add_heading(doc, "2. Architecture and Component Map", 1)
    add_para(doc, "The prototype is deliberately self-contained: HTML, CSS and vanilla JavaScript only. There is no Firebase, no server, no API key and no build step. This is appropriate for a 10 percent proof of concept because the AI workflow is visible and repeatable during marking.")
    add_table(
        doc,
        ["File", "Main object", "Important functions", "Responsibility"],
        [
            ["js/kb.js", "KB, KBUtil", "findStop, findStops, routesServing, routesBetween, nextDeparture, setClock, subscribe, setDelayed, logFeedback", "Knowledge base, route/stop lookup, directional route matching, timetable estimates, subscriptions, delays and feedback state."],
            ["js/intent.js", "Intent", "classify(query)", "Scores six major intents and exposes confidence in the UI."],
            ["js/retriever.js", "Retriever", "retrieve(query, k)", "RAG-style token retrieval over stops, routes, FAQs, alerts and source notes."],
            ["js/resolution.js", "Resolution", "prove(user, route)", "First-order resolution-refutation proof for NotifyUser(user, route, delay)."],
            ["js/chat.js", "Chat", "send(text), init()", "Conversation controller, session memory, response cards and AI pipeline updates."],
            ["js/app.js", "App", "showScreen, showStop, showRouteTimetable, renderAlerts, renderAdmin, runProof", "Screen routing, timetable, Alerts, Feedback, Staff Control Panel and proof display."],
        ],
        widths=[Inches(0.78), Inches(0.85), Inches(2.25), Inches(2.72)],
    )
    add_para(doc, "Runtime flow: user question -> Intent.classify() -> Retriever.retrieve() -> per-intent response handler -> grounded UI card + AI pipeline panel. Staff delay flow: Staff Control Panel toggle -> KBUtil.setDelayed() -> Resolution.prove() -> Alerts screen for subscribed users only.")

    add_heading(doc, "3. Data Source and Accuracy Policy", 1)
    add_para(doc, "Use this explanation during presentation: route names and stop sequences are aligned to public UTM/DVC/KDOJ sources where available. Timetable, ETA, delay duration, operating state and walking notes are configured estimates for demonstrating the AI concept, because no verified live UTM shuttle feed is connected.")
    add_table(
        doc,
        ["Data shown", "Status", "Administrator explanation"],
        [
            ["BAS A/B/C/D/E/F/G/H route labels and stop sequences", "Public-list aligned", "Used as route/sequence reference where public UTM/DVC/KDOJ listings are available. Do not claim they prove current live operation."],
            ["Stop labels such as CP, KP, K9/K10, KTC, KDOJ, KDSE, KTDI, N24, SKT, P19, T02, T08, V01", "Public stop labels", "Used as route-stop codes because they appear in public route listings."],
            ["FKE mapping", "Partially verified", "Mapped to P19 / FKE Area because public listings include P19 and FKE-area stop references. Production should confirm exact shelter."],
            ["FC mapping", "Prototype alias", "Mapped to N24 / Cluster Area for the report demo. Exact current FC shuttle-stop mapping was not found."],
            ["PSZ Library", "Landmark only", "The app can identify PSZ as a campus landmark, but it does not invent a direct shuttle route to PSZ."],
            ["Next bus, full-day timetable, service window, ETA and delay duration", "Prototype estimate", "Calculated from configured operating windows/headways and demo time. Not an official live timetable."],
            ["Delay alert", "Prototype estimate", "Staff toggle represents transport staff publishing a disruption so the logic proof and alert workflow can be shown."],
        ],
        widths=[Inches(2.05), Inches(1.1), Inches(3.45)],
    )
    add_para(doc, "Sources checked on 2 July 2026: UTM DVC Development shuttle bus schedule page, KDOJ UTM Bus Schedule route list, and the UTM JB 2025 shuttle timetable PDF.")

    add_heading(doc, "4. Screenshot-Guided Operating Procedure", 1)
    add_para(doc, "The following figures are current screenshots from the redesigned prototype. The arrows identify the control or evidence to point at during the presentation.")

    add_heading(doc, "4.1 Home Screen and Navigation", 2)
    add_steps(doc, [
        "Open index.html. Confirm the phone status clock is visible at the top right.",
        "Use Home cards or the bottom tab bar to start the main workflows.",
        "Keep the AI pipeline panel visible when explaining originality.",
    ])
    add_image(doc, "01-home-overview.png", "Figure 1. Home screen with demo clock, navigation cards and AI pipeline panel.")

    add_heading(doc, "4.2 Chat Schedule Answer", 2)
    add_steps(doc, [
        "Open Chat and ask: When is the next bus to FKE?",
        "Point to intent confidence to show the intent recognition result.",
        "Point to first bus, last bus, service window and frequency to show the app now behaves like a transport product.",
        "Explain the source boundary: public route data plus configured timetable estimate.",
    ])
    add_image(doc, "02-chat-schedule.png", "Figure 2. Schedule card with first bus, last bus, service window, frequency and AI pipeline evidence.")

    add_heading(doc, "4.3 Timetable Search and Full-Day Route Timetable", 2)
    add_steps(doc, [
        "Open Bus Timetable.",
        "Search by route, stop or faculty alias, for example FKE.",
        "Tap the BAS G route card to open its route timetable.",
        "Scroll to show the full-day departure groups and the highlighted next bus.",
    ])
    add_image(doc, "03-timetable-search.png", "Figure 3. Timetable search filtering routes by FKE alias.")
    add_image(doc, "04-route-timetable-detail.png", "Figure 4. Route timetable summary with route direction and key timing metrics.")
    add_image(doc, "04b-route-timetable-times.png", "Figure 5. Full-day departure grid with morning, afternoon and evening groups.")

    add_heading(doc, "4.4 Route Guidance", 2)
    add_steps(doc, [
        "Ask: How do I get from KTDI to P19 FKE?",
        "Show that the route engine uses the ordered stop sequence, so origin must appear before destination.",
        "Point to the AI panel showing route_guidance intent and retrieved route/stop chunks.",
    ])
    add_image(doc, "05-route-guidance.png", "Figure 6. Directional route guidance from KTDI to P19 / FKE Area.")

    add_heading(doc, "4.5 Bus Stop Detail", 2)
    add_steps(doc, [
        "Ask: Where is the FC bus stop?",
        "The app opens the Bus Stop Detail screen for N24 / Cluster Area.",
        "Point to Data status to show uncertainty is clearly described instead of overstated.",
    ])
    add_image(doc, "06-bus-stop-detail.png", "Figure 7. Bus Stop Detail with routes, data status, facilities and walking guidance.")

    add_heading(doc, "4.6 Alerts and Subscription Filtering", 2)
    add_steps(doc, [
        "In Staff Control Panel, delay a route that Ali is not subscribed to, such as BAS G.",
        "Open Alerts. The route delay is hidden because Ali is not subscribed to that route.",
        "Delay BAS A, which Ali is subscribed to. The alert becomes visible with scheduled time, ETA, service window and frequency.",
    ])
    add_image(doc, "07-alerts-subscription-filter.png", "Figure 8. Alerts screen proving unsubscribed route delays stay hidden.")
    add_image(doc, "08-alerts-after-proof.png", "Figure 9. Subscribed BAS A delay alert with ETA and route timing details.")

    add_heading(doc, "4.7 Feedback and Staff Report Follow-Up", 2)
    add_steps(doc, [
        "Open Feedback / Escalate.",
        "Select an issue type, enter details and submit the report.",
        "Open Staff Control Panel. The same report appears under Reported issues for staff follow-up.",
    ])
    add_image(doc, "09-feedback-submission.png", "Figure 10. Feedback report logged by the passenger.")
    add_image(doc, "10-staff-report-feed.png", "Figure 11. Staff Control Panel showing the submitted passenger report.")

    add_heading(doc, "4.8 Staff Delay Proof", 2)
    add_para(doc, "The Staff Control Panel is inside the prototype so the marker can trigger repeatable proof scenarios. The Word manual remains the separate admin deliverable.")
    add_steps(doc, [
        "Open Staff Control Panel.",
        "Toggle BAS A1/A2 to delayed.",
        "The app runs Resolution.prove('ali', 'route_a').",
        "The proof panel derives the empty clause, proving NotifyUser(ali, route_a, delay).",
        "Open Alerts to verify only the subscribed route notification is visible.",
    ])
    add_table(
        doc,
        ["Clause", "Meaning"],
        [
            ["P1 WantsRouteAlert(ali, route_a)", "Ali subscribes to BAS A route alerts."],
            ["P2 Delayed(route_a)", "BAS A is currently marked delayed by staff."],
            ["P3 not Delayed(r) OR NeedDelayNotification(r)", "A delayed route requires a delay notification."],
            ["P4 not WantsRouteAlert(u,r) OR not NeedDelayNotification(r) OR NotifyUser(u,r,delay)", "A subscribed user is notified when a notification is needed."],
            ["Negated goal: not NotifyUser(ali, route_a, delay)", "Resolution assumes the opposite and derives a contradiction."],
        ],
        widths=[Inches(2.45), Inches(4.15)],
    )
    add_image(doc, "11-resolution-proof.png", "Figure 12. Staff proof trace after BAS A is marked delayed.")
    add_image(doc, "11b-resolution-proof-box.png", "Figure 13. Cropped proof trace showing the empty clause.", width=3.15)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "5. Configuration Guide", 1)
    add_heading(doc, "5.1 Edit Stops", 2)
    add_para(doc, "Stops are configured in KB.stops inside js/kb.js. Use aliases for names students actually type, and always include a status field that states whether the mapping is public, verified, or prototype-only.")
    add_code(doc, '{ id: "p19", name: "P19 / FKE Area", aliases: ["p19", "fke", "electrical"], status: "public stop/area label with faculty alias", facilities: ["Faculty area", "Lecture halls"], walking: "Alight at P19 for the FKE-area prototype. Confirm the exact shelter in production." }')

    add_heading(doc, "5.2 Edit Routes and Timetables", 2)
    add_para(doc, "Routes are configured in KB.routes. The full-day timetable is generated from operating.start, operating.end and headway. This keeps the demo repeatable without pretending to use live GPS.")
    add_code(doc, '{ id: "route_g", name: "BAS G1/G2/G3 - KTR/KTHO/KTDI to SKT", stops: ["ktr","ktho","ktdi","n24","skt","p19","cp"], operating: { start: "07:00", end: "18:30" }, headway: 20, source: "Public UTM/DVC/KDOJ route listing; timetable estimate configured in app." }')

    add_heading(doc, "5.3 Tune Intents and Retrieval", 2)
    add_bullets(doc, [
        "Edit CUES in js/intent.js when common student phrases are missing.",
        "Strong cues score 2; weak cues score 1. Avoid using overly common words as strong cues.",
        "Edit KB.faqs and route/stop aliases to improve RAG-style retrieval results.",
        "After changing intent cues, test schedule, route, arrival, bus-stop, alerts, feedback and FAQ queries.",
    ])

    add_heading(doc, "5.4 Configure Demo Time", 2)
    add_para(doc, "The default demo time is 10:00 so presentation does not depend on the real current time. Staff can switch to 08:00, 10:00, 13:00, 17:30, 20:30 or live device time from Staff Control Panel.")

    add_heading(doc, "5.5 Configure Subscriptions and Delay Proof", 2)
    add_bullets(doc, [
        "KB.subscriptions stores WantsRouteAlert(user, route) facts. The demo user is ali.",
        "KB.delayedRoutes stores Delayed(route) facts. Staff toggles update this list.",
        "Resolution.prove(user, route) only proves NotifyUser when both subscription and delay facts exist.",
        "If a delayed route is not subscribed, Alerts correctly hides the notification for Ali.",
    ])

    add_heading(doc, "6. Troubleshooting", 1)
    add_table(
        doc,
        ["Symptom", "Likely cause", "Fix"],
        [
            ["Blank screen", "Script load issue or missing file", "Reload and ensure kb.js, intent.js, retriever.js, resolution.js, chat.js and app.js are loaded in index.html order."],
            ["Schedule says closed now", "Demo/live clock outside service window", "Staff Control Panel -> Demo time -> 10:00 for normal presentation."],
            ["Wrong route result", "No directional route sequence serves origin before destination", "Use CP as interchange or add a verified route sequence. Do not invent official data."],
            ["No delay alert appears", "Ali is not subscribed to the delayed route", "Turn on the route subscription or use BAS A, which is seeded for the proof."],
            ["Proof does not run", "Route is not marked delayed or user lacks subscription", "Toggle BAS A delayed and confirm KB.subscriptions contains ali/route_a."],
            ["Feedback not visible to staff", "Staff Control Panel has not been opened/refreshed after submission", "Open Staff Control Panel; renderStaffIssues() reads KB.feedbackLog."],
            ["Teacher asks if times are official", "Prototype uses configured estimates", "State clearly: route names/sequences are public-list aligned; timing/ETA/delay are POC estimates."],
        ],
        widths=[Inches(1.45), Inches(2.1), Inches(3.05)],
        size=7.6,
    )

    add_heading(doc, "7. Demo Checklist", 1)
    add_bullets(doc, [
        "Open index.html or serve the prototype folder with python3 -m http.server 8000.",
        "Set Staff Control Panel -> Demo time to 10:00 before the main demo.",
        "Show Home, Chat schedule, Timetable search, route timetable, route guidance, bus-stop detail, Alerts, Feedback and Staff Control Panel.",
        "Submit a Feedback report and verify it appears under Staff Control Panel -> Reported issues.",
        "Toggle BAS A delayed, show the resolution proof, then open Alerts to show the subscribed notification.",
        "Say the data boundary clearly: public route names/sequences where available; timetable, ETA and delay values are configured POC estimates.",
    ])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("TransitAI UTM Prototype | Administrator Manual | Group 5")
    set_font(fr, size=8)

    doc.save(DOCX_OUT)


def build_html_and_markdown():
    html_parts = [
        "<!DOCTYPE html>",
        '<html lang="en"><head><meta charset="UTF-8" />',
        '<meta name="viewport" content="width=device-width, initial-scale=1.0" />',
        "<title>TransitAI UTM - Administrator Manual</title>",
        "<style>",
        "body{font-family:Arial,Helvetica,sans-serif;color:#000;max-width:980px;margin:0 auto;padding:32px;line-height:1.45}",
        "h1,h2,h3{color:#000} h1{text-align:center} table{border-collapse:collapse;width:100%;margin:12px 0;font-size:13px}",
        "th,td{border:1px solid #000;padding:7px;vertical-align:top;text-align:left;background:#fff;color:#000}",
        "pre,code{font-family:'Courier New',monospace;color:#000} pre{border:1px solid #000;padding:10px;white-space:pre-wrap}",
        "figure{margin:18px 0;text-align:center;break-inside:avoid} figure img{max-width:100%;border:1px solid #000}",
        "figcaption{font-size:12px;margin-top:6px;color:#000;font-style:italic}.title-note{text-align:center}",
        "@media print{body{padding:0}figure{page-break-inside:avoid}}",
        "</style></head><body>",
        "<h1>TransitAI UTM<br>Administrator Manual</h1>",
        '<p class="title-note">Web-Based Campus Transport Chatbot Prototype<br>MECS0033-52 Artificial Intelligence, Section 52, Group 5</p>',
    ]
    md_parts = [
        "# TransitAI UTM - Administrator Manual",
        "",
        "Web-Based Campus Transport Chatbot Prototype",
        "MECS0033-52 Artificial Intelligence, Section 52, Group 5",
        "",
    ]

    for event in manual_events:
        kind = event[0]
        if kind == "heading":
            _, level, text = event
            html_parts.append(f"<h{level + 1}>{h(text)}</h{level + 1}>")
            md_parts.extend([f"{'#' * (level + 1)} {text}", ""])
        elif kind == "p":
            html_parts.append(f"<p>{h(event[1])}</p>")
            md_parts.extend([event[1], ""])
        elif kind == "bullets":
            html_parts.append("<ul>")
            for item in event[1]:
                html_parts.append(f"<li>{h(item)}</li>")
            html_parts.append("</ul>")
            md_parts.extend([f"- {item}" for item in event[1]] + [""])
        elif kind == "steps":
            html_parts.append("<ol>")
            for item in event[1]:
                html_parts.append(f"<li>{h(item)}</li>")
            html_parts.append("</ol>")
            md_parts.extend([f"{idx}. {item}" for idx, item in enumerate(event[1], 1)] + [""])
        elif kind == "code":
            html_parts.append(f"<pre>{h(event[1])}</pre>")
            md_parts.extend(["```js", event[1], "```", ""])
        elif kind == "table":
            _, headers, rows = event
            html_parts.append("<table><thead><tr>")
            for header in headers:
                html_parts.append(f"<th>{h(header)}</th>")
            html_parts.append("</tr></thead><tbody>")
            for row in rows:
                html_parts.append("<tr>")
                for cell in row:
                    html_parts.append(f"<td>{h(cell)}</td>")
                html_parts.append("</tr>")
            html_parts.append("</tbody></table>")
            md_parts.append("| " + " | ".join(headers) + " |")
            md_parts.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in rows:
                md_parts.append("| " + " | ".join(str(cell).replace("|", "/") for cell in row) + " |")
            md_parts.append("")
        elif kind == "image":
            _, file_name, caption = event
            src = f"manual_screenshots_annotated/{file_name}"
            html_parts.append(f'<figure><img src="{h(src)}" alt="{h(caption)}" /><figcaption>{h(caption)}</figcaption></figure>')
            md_parts.extend([f"![{caption}]({src})", f"*{caption}*", ""])

    html_parts.append("<footer><p>TransitAI UTM Prototype | Administrator Manual | Group 5</p></footer>")
    html_parts.append("</body></html>")
    HTML_OUT.write_text("\n".join(html_parts), encoding="utf-8")
    MD_OUT.write_text("\n".join(md_parts), encoding="utf-8")


build_docx()
build_html_and_markdown()
print(DOCX_OUT)
print(HTML_OUT)
print(MD_OUT)
